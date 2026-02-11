"""Application Composer Agent - positioning strategy, resume/cover letter generation."""

import json
import re
from datetime import datetime, timezone
from pathlib import Path

from rich.console import Console
from rich.panel import Panel
from rich.progress import Progress, SpinnerColumn, TextColumn

from data_store import DataStore
from skills import (
    JobDescriptionAnalyzerSkill,
    ResumeGeneratorSkill,
    CoverLetterGeneratorSkill,
    InterviewPrepSkill,
    SkillContext,
)
from .base_agent import BaseAgent

console = Console()


class ApplicationComposerAgent(BaseAgent):
    """Agent that handles application materials - analysis, resumes, cover letters.

    Consolidates functionality from JobResearcherAgent:
    - analyze_job
    - generate_resume
    - generate_cover_letter
    - improve_resume
    - regenerate_output (+ regenerate_pdf backward-compat wrapper)
    """

    def __init__(self, config: dict):
        """Initialize the application composer agent.

        Args:
            config: Configuration dictionary.
        """
        super().__init__(config)
        self.data_store = DataStore(config)

        # Initialize skills
        self.job_analyzer = JobDescriptionAnalyzerSkill(
            self.client, self.data_store, config
        )
        self.resume_generator = ResumeGeneratorSkill(
            self.client, self.data_store, config
        )
        self.cover_letter_generator = CoverLetterGeneratorSkill(
            self.client, self.data_store, config
        )
        self.interview_prep = InterviewPrepSkill(
            self.client, self.data_store, config
        )

        # Ensure output directories exist
        (self.output_dir / "resumes").mkdir(parents=True, exist_ok=True)
        (self.output_dir / "cover-letters").mkdir(parents=True, exist_ok=True)
        (self.output_dir / "analysis").mkdir(parents=True, exist_ok=True)
        (self.output_dir / "interview-prep").mkdir(parents=True, exist_ok=True)

    # =========================================================================
    # Job Analysis
    # =========================================================================

    def analyze_job(self, job_id: str) -> dict | None:
        """Analyze a job posting and match against user's profile.

        Args:
            job_id: ID of the job to analyze.

        Returns:
            Analysis dictionary or None if failed.
        """
        # Find the job
        job = self.data_store.get_job(job_id)
        if not job:
            console.print(f"[red]Job not found: {job_id}[/red]")
            return None

        # Load base resume
        resume_text = self._load_base_resume()
        if not resume_text:
            console.print("[red]Could not load base resume[/red]")
            return None

        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            console=console,
        ) as progress:
            task = progress.add_task("Analyzing job requirements...", total=None)

            # Build skill context
            context = SkillContext(config=self.config)

            # Execute job analyzer skill
            result = self.job_analyzer.execute(context, job, resume_text)

            if not result.success:
                console.print(f"[red]Analysis failed: {result.error}[/red]")
                return None

            analysis_result = result.data
            raw_analysis = result.metadata.get("raw_analysis", {})

            progress.update(task, description="Saving analysis...")

            # Save analysis
            self._save_analysis(job_id, job, raw_analysis)

        # Print summary
        self._print_analysis(job, analysis_result)

        return raw_analysis

    # =========================================================================
    # Resume Generation
    # =========================================================================

    def generate_resume(self, job_id: str, output_format: str = "pdf") -> str | None:
        """Generate a customized resume for a job.

        Args:
            job_id: ID of the job to generate resume for.

        Returns:
            Path to saved resume markdown, or None if failed.
        """
        # Find the job
        job = self.data_store.get_job(job_id)
        if not job:
            console.print(f"[red]Job not found: {job_id}[/red]")
            return None

        # Determine role lens
        role_lens = self.job_analyzer.determine_role_lens(job)
        console.print(f"[dim]Role lens: {role_lens}[/dim]")

        # Load base resume
        resume_text = self._load_base_resume()
        if not resume_text:
            console.print("[red]Could not load base resume[/red]")
            return None

        # Load analysis if exists
        analysis = self._load_analysis(job_id)

        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            console=console,
        ) as progress:
            task = progress.add_task("Generating customized resume...", total=None)

            # Build skill context
            context = SkillContext(config=self.config)

            # Execute resume generator skill
            result = self.resume_generator.execute(
                context,
                job=job,
                base_resume=resume_text,
                analysis=analysis,
                role_lens=role_lens,
            )

            if not result.success:
                console.print(f"[red]Resume generation failed: {result.error}[/red]")
                return None

            resume_md = result.data.resume_markdown

            progress.update(task, description="Saving resume...")

            # Save markdown
            company_name = self._sanitize_filename(job.get("company", "Unknown"))
            job_title = self._sanitize_filename(job.get("title", "Unknown"))
            md_path = self.output_dir / "resumes" / f"Ron Lin Resume - {company_name} - {job_title}.md"

            with open(md_path, "w") as f:
                f.write(resume_md)

            # Generate output format(s)
            fmt_label = "PDF + DOCX" if output_format == "both" else output_format.upper()
            progress.update(task, description=f"Generating {fmt_label}...")
            output_paths = self._convert_document(md_path, "resume", output_format)

        console.print(f"[green]Resume saved to:[/green] {md_path}")
        for fmt, path in output_paths.items():
            if path:
                console.print(f"[green]{fmt.upper()} saved to:[/green] {path}")

        return str(md_path)

    def improve_resume(self, job_id: str, output_format: str = "pdf") -> str | None:
        """Improve an existing resume using three-phase pipeline.

        Phase 1: Generate edit plan (3-8 surgical edits) using analysis data
        Phase 2: Apply edits programmatically to preserve format
        Phase 3: Credibility audit on changed lines only

        Args:
            job_id: ID of the job the resume is for.

        Returns:
            Path to improved resume markdown, or None if failed.
        """
        # Find the job
        job = self.data_store.get_job(job_id)
        if not job:
            console.print(f"[red]Job not found: {job_id}[/red]")
            return None

        # Find existing resume for this job
        resume_path = self.find_document_by_job_id(job_id, "resume")
        if not resume_path:
            console.print(f"[red]No existing resume found for job: {job_id}[/red]")
            console.print("[dim]Run 'scout resume <job_id>' first to generate a resume.[/dim]")
            return None

        # Load the current resume
        with open(resume_path) as f:
            current_resume = f.read()

        # Load base resume for reference
        base_resume = self._load_base_resume()
        if not base_resume:
            console.print("[red]Could not load base resume[/red]")
            return None

        # Load analysis — auto-run analyze if missing
        analysis = self._load_analysis(job_id)
        if not analysis:
            console.print("[yellow]No analysis found. Running analyze first...[/yellow]\n")
            analysis_result = self.analyze_job(job_id)
            if analysis_result:
                analysis = analysis_result
            else:
                console.print("[yellow]Analysis failed — proceeding with degraded improvement[/yellow]\n")

        # Determine role lens
        role_lens = self.job_analyzer.determine_role_lens(job)

        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            console=console,
        ) as progress:
            # Phase 1: Generate edit plan
            task = progress.add_task("Phase 1: Generating edit plan...", total=None)

            context = SkillContext(config=self.config)

            edit_plan_result = self.resume_generator.plan_resume_edits(
                context,
                job=job,
                current_resume=current_resume,
                base_resume=base_resume,
                analysis=analysis,
                role_lens=role_lens,
            )

            if not edit_plan_result.success:
                console.print(f"[red]Edit plan generation failed: {edit_plan_result.error}[/red]")
                return None

            edit_plan = edit_plan_result.data

            # Save versioned edit plan to disk
            edit_plan_path = self._save_edit_plan(job_id, edit_plan)

            # Phase 2: Apply edits programmatically
            progress.update(task, description="Phase 2: Applying edits...")

            apply_result = self.resume_generator.apply_resume_edits(
                current_resume, edit_plan
            )

            if not apply_result.success:
                console.print(f"[red]Edit application failed: {apply_result.error}[/red]")
                return None

            modified_resume = apply_result.data["resume"]
            apply_report = apply_result.data["report"]

            # Phase 3: Credibility audit
            progress.update(task, description="Phase 3: Credibility audit...")

            audit_result = self.resume_generator.audit_resume_edits(
                context,
                modified_resume=modified_resume,
                original_resume=current_resume,
                base_resume=base_resume,
                job=job,
                edit_plan=edit_plan,
            )

            if audit_result.success:
                final_resume = audit_result.data["resume"]
                audit_report = audit_result.data["report"]
            else:
                # Audit failed — use the Phase 2 output as-is
                console.print(f"[yellow]Credibility audit failed, using Phase 2 output: {audit_result.error}[/yellow]")
                final_resume = modified_resume
                audit_report = []

            progress.update(task, description="Saving improved resume...")

            # Save the improved resume
            with open(resume_path, "w") as f:
                f.write(final_resume)

            fmt_label = "PDF + DOCX" if output_format == "both" else output_format.upper()
            progress.update(task, description=f"Generating {fmt_label}...")

            # Regenerate output format(s)
            output_paths = self._convert_document(resume_path, "resume", output_format)

        # Print edit plan summary
        self._print_edit_summary(edit_plan, apply_report, audit_report)

        console.print(f"\n[green]Resume improved and saved to:[/green] {resume_path}")
        if edit_plan_path:
            console.print(f"[dim]Edit plan saved to:[/dim] {edit_plan_path}")
        for fmt, path in output_paths.items():
            if path:
                console.print(f"[green]{fmt.upper()} saved to:[/green] {path}")

        return str(resume_path)

    def _save_edit_plan(self, job_id: str, edit_plan: dict) -> Path | None:
        """Save versioned edit plan to disk.

        Finds the next version number and saves as {job-id}-edit-plan-v{N}.json.
        """
        analysis_dir = self.output_dir / "analysis"
        analysis_dir.mkdir(parents=True, exist_ok=True)

        # Find next version number
        version = 1
        while (analysis_dir / f"{job_id}-edit-plan-v{version}.json").exists():
            version += 1

        edit_plan_path = analysis_dir / f"{job_id}-edit-plan-v{version}.json"

        data = {
            "job_id": job_id,
            "version": version,
            "edit_plan": edit_plan,
            "created_at": datetime.now(timezone.utc).isoformat(),
        }

        with open(edit_plan_path, "w") as f:
            json.dump(data, f, indent=2)

        return edit_plan_path

    def _print_edit_summary(
        self, edit_plan: dict, apply_report: list[dict], audit_report: list[dict]
    ) -> None:
        """Print structured edit summary using Rich panels."""
        edits = edit_plan.get("edit_plan", [])

        if not edits:
            console.print("[yellow]No edits were proposed.[/yellow]")
            return

        # Edit plan summary
        edit_lines = []
        for i, edit in enumerate(edits, 1):
            edit_type = edit.get("edit_type", "replace")
            target = edit.get("target", "unknown")
            rationale = edit.get("rationale", "")
            edit_lines.append(f"[bold]{i}. [{edit_type.upper()}][/bold] {target}")
            edit_lines.append(f"   [dim]{rationale}[/dim]")

        console.print(Panel(
            "\n".join(edit_lines),
            title=f"[bold green]Edit Plan ({len(edits)} edits)[/bold green]",
            border_style="green",
        ))

        # Apply report — show any failures
        failed = [r for r in apply_report if not r.get("applied")]
        if failed:
            fail_lines = [f"- {r.get('target', '?')}: {r.get('reason', 'unknown')}" for r in failed]
            console.print(Panel(
                "\n".join(fail_lines),
                title="[bold yellow]Edits That Needed Fallback[/bold yellow]",
                border_style="yellow",
            ))

        # Audit report — show any credibility flags
        if audit_report:
            audit_lines = [f"- {item}" for item in audit_report]
            console.print(Panel(
                "\n".join(audit_lines),
                title="[bold cyan]Credibility Audit[/bold cyan]",
                border_style="cyan",
            ))

        # Remaining gaps
        remaining_gaps = edit_plan.get("remaining_gaps", [])
        if remaining_gaps:
            gap_lines = [f"- {gap}" for gap in remaining_gaps]
            console.print(Panel(
                "\n".join(gap_lines),
                title="[dim]Remaining Gaps (cannot address without fabrication)[/dim]",
                border_style="dim",
            ))

    # =========================================================================
    # Cover Letter Generation
    # =========================================================================

    def generate_cover_letter(self, job_id: str, output_format: str = "pdf") -> str | None:
        """Generate a cover letter for a job.

        Args:
            job_id: ID of the job to generate cover letter for.

        Returns:
            Path to saved cover letter markdown, or None if failed.
        """
        # Find the job
        job = self.data_store.get_job(job_id)
        if not job:
            console.print(f"[red]Job not found: {job_id}[/red]")
            return None

        # Determine role lens
        role_lens = self.job_analyzer.determine_role_lens(job)
        console.print(f"[dim]Role lens: {role_lens}[/dim]")

        # Load base resume
        resume_text = self._load_base_resume()
        if not resume_text:
            console.print("[red]Could not load base resume[/red]")
            return None

        # Load analysis if exists
        analysis = self._load_analysis(job_id)

        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            console=console,
        ) as progress:
            task = progress.add_task("Generating cover letter...", total=None)

            # Build skill context
            context = SkillContext(config=self.config)

            # Execute cover letter generator skill
            result = self.cover_letter_generator.execute(
                context,
                job=job,
                base_resume=resume_text,
                analysis=analysis,
                role_lens=role_lens,
            )

            if not result.success:
                console.print(f"[red]Cover letter generation failed: {result.error}[/red]")
                return None

            cover_letter_md = result.data.cover_letter_markdown

            progress.update(task, description="Saving cover letter...")

            # Save markdown
            company_name = self._sanitize_filename(job.get("company", "Unknown"))
            job_title = self._sanitize_filename(job.get("title", "Unknown"))
            md_path = self.output_dir / "cover-letters" / f"Cover Letter - {company_name} - {job_title}.md"

            with open(md_path, "w") as f:
                f.write(cover_letter_md)

            # Generate output format(s)
            fmt_label = "PDF + DOCX" if output_format == "both" else output_format.upper()
            progress.update(task, description=f"Generating {fmt_label}...")
            output_paths = self._convert_document(md_path, "cover-letter", output_format)

        console.print(f"[green]Cover letter saved to:[/green] {md_path}")
        for fmt, path in output_paths.items():
            if path:
                console.print(f"[green]{fmt.upper()} saved to:[/green] {path}")

        return str(md_path)

    # =========================================================================
    # Interview Prep
    # =========================================================================

    def generate_interview_prep(self, job_id: str) -> str | None:
        """Generate screening interview preparation materials for a job.

        Args:
            job_id: ID of the job to prepare for.

        Returns:
            Path to saved interview prep markdown, or None if failed.
        """
        # Find the job
        job = self.data_store.get_job(job_id)
        if not job:
            console.print(f"[red]Job not found: {job_id}[/red]")
            return None

        # Determine role lens
        role_lens = self.job_analyzer.determine_role_lens(job)
        console.print(f"[dim]Role lens: {role_lens}[/dim]")

        # Load base resume
        resume_text = self._load_base_resume()
        if not resume_text:
            console.print("[red]Could not load base resume[/red]")
            return None

        # Load analysis — auto-run analyze if missing
        analysis = self._load_analysis(job_id)
        if not analysis:
            console.print("[yellow]No analysis found. Running analyze first...[/yellow]\n")
            analysis_result = self.analyze_job(job_id)
            if analysis_result:
                analysis = analysis_result
            else:
                console.print("[yellow]Analysis failed — proceeding without analysis data[/yellow]\n")

        # Load company research (optional — enriches "Why Company?" section)
        company_name = job.get("company", "")
        company_research = None
        if company_name:
            company_slug = re.sub(r"[^a-z0-9]+", "-", company_name.lower()).strip("-")
            company_research = self.data_store.get_research(company_slug)
            if company_research:
                console.print(f"[dim]Using company research for {company_name}[/dim]")

        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            console=console,
        ) as progress:
            task = progress.add_task("Generating interview prep...", total=None)

            # Build skill context
            context = SkillContext(config=self.config)

            # Execute interview prep skill
            result = self.interview_prep.execute(
                context,
                job=job,
                base_resume=resume_text,
                analysis=analysis,
                role_lens=role_lens,
                company_research=company_research,
            )

            if not result.success:
                console.print(f"[red]Interview prep generation failed: {result.error}[/red]")
                return None

            prep_result = result.data
            prep_md = prep_result.prep_markdown

            progress.update(task, description="Saving interview prep...")

            # Save markdown
            sanitized_company = self._sanitize_filename(job.get("company", "Unknown"))
            sanitized_title = self._sanitize_filename(job.get("title", "Unknown"))
            md_path = self.output_dir / "interview-prep" / f"Interview Prep - {sanitized_company} - {sanitized_title}.md"

            with open(md_path, "w") as f:
                f.write(prep_md)

        # Print summary
        self._print_interview_prep_summary(prep_result, md_path)

        return str(md_path)

    def _print_interview_prep_summary(self, prep_result, md_path: Path) -> None:
        """Print interview prep summary to console."""
        prep_md = prep_result.prep_markdown

        # Extract elevator pitch preview (text between "## 1." and the next "##")
        pitch_preview = ""
        lines = prep_md.split("\n")
        in_pitch = False
        pitch_lines = []
        for line in lines:
            if "Elevator Pitch" in line and line.startswith("#"):
                in_pitch = True
                continue
            if in_pitch and line.startswith("## "):
                break
            if in_pitch and line.strip() and not line.startswith("#"):
                pitch_lines.append(line.strip())

        if pitch_lines:
            # Show first 3 sentences of the pitch
            pitch_text = " ".join(pitch_lines)
            sentences = pitch_text.replace(". ", ".|").split("|")
            pitch_preview = " ".join(sentences[:3]).strip()
            if not pitch_preview.endswith("."):
                pitch_preview += "..."

        if pitch_preview:
            console.print(Panel(
                pitch_preview,
                title="[bold green]Elevator Pitch Preview[/bold green]",
                border_style="green",
            ))

        # Stats
        dc_count = prep_result.domain_connection_count
        if dc_count > 0:
            console.print(f"  Domain connection talking points: [cyan]{dc_count}[/cyan]")
        console.print(f"  Sections generated: [cyan]{prep_result.section_count}[/cyan]")

        console.print(f"\n[green]Interview prep saved to:[/green] {md_path}")

    # =========================================================================
    # Document Management
    # =========================================================================

    def find_document_by_job_id(self, job_id: str, doc_type: str) -> Path | None:
        """Find an existing markdown document by job_id.

        Args:
            job_id: ID of the job.
            doc_type: Type of document ("resume" or "cover-letter").

        Returns:
            Path to the document or None if not found.
        """
        # Find the job to get company and title
        job = self.data_store.get_job(job_id)
        if not job:
            return None

        company_name = self._sanitize_filename(job.get("company", "Unknown"))
        job_title = self._sanitize_filename(job.get("title", "Unknown"))

        if doc_type == "resume":
            search_dir = self.output_dir / "resumes"
            pattern = f"Ron Lin Resume - {company_name} - {job_title}.md"
        else:
            search_dir = self.output_dir / "cover-letters"
            pattern = f"Cover Letter - {company_name} - {job_title}.md"

        # Try exact match first
        exact_path = search_dir / pattern
        if exact_path.exists():
            return exact_path

        # Fall back to searching for partial match
        if search_dir.exists():
            for md_file in search_dir.glob("*.md"):
                if company_name in md_file.name and job_title[:20] in md_file.name:
                    return md_file

        return None

    def regenerate_output(self, md_path: Path, doc_type: str, output_format: str = "pdf") -> dict:
        """Regenerate output file(s) from an existing markdown file.

        Args:
            md_path: Path to the markdown file.
            doc_type: Type of document ("resume" or "cover-letter").
            output_format: "pdf", "docx", or "both".

        Returns:
            Dict with "pdf" and/or "docx" keys mapping to output Paths (or None if failed).
        """
        return self._convert_document(md_path, doc_type, output_format)

    def regenerate_pdf(self, md_path: Path, doc_type: str) -> Path | None:
        """Regenerate PDF from an existing markdown file.

        Backward-compatible wrapper around regenerate_output().

        Args:
            md_path: Path to the markdown file.
            doc_type: Type of document ("resume" or "cover-letter").

        Returns:
            Path to the PDF or None if failed.
        """
        results = self.regenerate_output(md_path, doc_type, "pdf")
        return results.get("pdf")

    # =========================================================================
    # Private Methods
    # =========================================================================

    def _load_base_resume(self) -> str | None:
        """Load the base resume from markdown file."""
        resume_path = self.input_dir / "base-resume.md"

        if not resume_path.exists():
            console.print(f"[yellow]Resume not found at {resume_path}[/yellow]")
            return None

        try:
            return resume_path.read_text()
        except Exception as e:
            console.print(f"[red]Error reading resume: {e}[/red]")
            return None

    def _save_analysis(self, job_id: str, job: dict, analysis: dict) -> None:
        """Save job analysis to file."""
        output_path = self.output_dir / "analysis" / f"{job_id}-analysis.json"

        data = {
            "job_id": job_id,
            "job": job,
            "analysis": analysis,
            "analyzed_at": datetime.now(timezone.utc).isoformat(),
        }

        with open(output_path, "w") as f:
            json.dump(data, f, indent=2)

    def _load_analysis(self, job_id: str) -> dict | None:
        """Load existing analysis if available."""
        analysis_path = self.output_dir / "analysis" / f"{job_id}-analysis.json"

        if analysis_path.exists():
            with open(analysis_path) as f:
                data = json.load(f)
                return data.get("analysis")

        return None

    def _convert_document(self, md_path: Path, doc_type: str, output_format: str = "pdf") -> dict:
        """Convert markdown to the requested output format(s).

        Dispatches to _convert_to_pdf() and/or _convert_to_docx() based on
        the output_format parameter.

        Args:
            md_path: Path to the source markdown file.
            doc_type: Type of document ("resume" or "cover-letter").
            output_format: "pdf", "docx", or "both".

        Returns:
            Dict with "pdf" and/or "docx" keys mapping to output Paths (or None if failed).
        """
        results = {}

        if output_format in ("pdf", "both"):
            results["pdf"] = self._convert_to_pdf(md_path, doc_type)

        if output_format in ("docx", "both"):
            results["docx"] = self._convert_to_docx(md_path, doc_type)

        return results

    def _convert_to_pdf(self, md_path: Path, doc_type: str) -> Path | None:
        """Convert markdown to PDF using weasyprint."""
        try:
            import markdown
            from weasyprint import HTML

            # Read markdown
            with open(md_path) as f:
                md_content = f.read()

            # Normalize bullet characters
            md_content = re.sub(r'^•\s*', '- ', md_content, flags=re.MULTILINE)

            # Convert to HTML
            html_content = markdown.markdown(md_content, extensions=['tables'])

            # Load template based on document type
            template_dir = Path(__file__).parent.parent / "templates"
            template_file = "resume.html" if doc_type == "resume" else "cover-letter.html"
            template_path = template_dir / template_file

            with open(template_path) as f:
                template = f.read()

            # Insert content into template
            full_html = template.replace("{{content}}", html_content)

            # Generate PDF
            pdf_path = md_path.with_suffix('.pdf')
            HTML(string=full_html).write_pdf(pdf_path)

            return pdf_path

        except ImportError as e:
            console.print(f"[yellow]PDF generation skipped (missing dependency): {e}[/yellow]")
            return None
        except Exception as e:
            console.print(f"[yellow]PDF generation failed: {e}[/yellow]")
            return None

    def _convert_to_docx(self, md_path: Path, doc_type: str) -> Path | None:
        """Convert markdown to DOCX using python-docx.

        Produces an ATS-optimized DOCX file using Word's built-in styles
        for maximum compatibility with applicant tracking systems.

        Args:
            md_path: Path to the source markdown file.
            doc_type: Type of document ("resume" or "cover-letter").

        Returns:
            Path to the generated DOCX file, or None if generation failed.
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            md_content = md_path.read_text()

            # Normalize bullet characters
            md_content = re.sub(r'^•\s*', '- ', md_content, flags=re.MULTILINE)

            doc = Document()

            # Page setup — match PDF template margins
            for section in doc.sections:
                section.top_margin = Inches(0.4)
                section.bottom_margin = Inches(0.4)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)

            # Configure default font — match PDF: Helvetica Neue/Arial, 10pt, 1.3 line height
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(10)
            font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            style.paragraph_format.line_spacing = 1.3
            style.paragraph_format.space_before = Pt(3)
            style.paragraph_format.space_after = Pt(3)

            # H1: Name/title — 16pt, no top margin
            h1_style = doc.styles['Heading 1']
            h1_style.font.name = 'Arial'
            h1_style.font.size = Pt(16)
            h1_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            h1_style.font.bold = True
            h1_style.paragraph_format.space_before = Pt(0)
            h1_style.paragraph_format.space_after = Pt(2)

            # H2: Section headers — 11pt, uppercase, bottom border (matches PDF CSS)
            h2_style = doc.styles['Heading 2']
            h2_style.font.name = 'Arial'
            h2_style.font.size = Pt(11)
            h2_style.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
            h2_style.font.bold = True
            h2_style.paragraph_format.space_before = Pt(12)
            h2_style.paragraph_format.space_after = Pt(4)
            # Bottom border: border-bottom: 1px solid #999
            h2_pPr = h2_style.element.get_or_add_pPr()
            h2_pBdr = OxmlElement('w:pBdr')
            h2_bottom = OxmlElement('w:bottom')
            h2_bottom.set(qn('w:val'), 'single')
            h2_bottom.set(qn('w:sz'), '4')
            h2_bottom.set(qn('w:space'), '1')
            h2_bottom.set(qn('w:color'), '999999')
            h2_pBdr.append(h2_bottom)
            h2_pPr.append(h2_pBdr)

            # H3: Company/role — 10pt
            h3_style = doc.styles['Heading 3']
            h3_style.font.name = 'Arial'
            h3_style.font.size = Pt(10)
            h3_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            h3_style.font.bold = False
            h3_style.paragraph_format.space_before = Pt(12)
            h3_style.paragraph_format.space_after = Pt(2)

            # Configure list bullet style spacing
            if 'List Bullet' in doc.styles:
                lb_style = doc.styles['List Bullet']
                lb_style.font.name = 'Arial'
                lb_style.font.size = Pt(10)
                lb_style.paragraph_format.space_before = Pt(1)
                lb_style.paragraph_format.space_after = Pt(1)

            self._parse_markdown_to_docx(doc, md_content, doc_type)

            docx_path = md_path.with_suffix('.docx')
            doc.save(str(docx_path))
            return docx_path

        except ImportError as e:
            console.print(f"[yellow]DOCX generation skipped (missing dependency): {e}[/yellow]")
            return None
        except Exception as e:
            console.print(f"[yellow]DOCX generation failed: {e}[/yellow]")
            return None

    def _parse_markdown_to_docx(self, doc, md_content: str, doc_type: str) -> None:
        """Parse markdown content and build DOCX document elements.

        Handles the known resume/cover-letter markdown structure:
        - # H1 headings (name)
        - ## H2 headings (section headers)
        - ### H3 headings (company/role)
        - Bullet lists (- prefixed)
        - Bold (**text**) and italic (*text*) inline formatting
        - Horizontal rules (---)
        - Plain paragraphs

        Args:
            doc: python-docx Document instance.
            md_content: Raw markdown string.
            doc_type: "resume" or "cover-letter".
        """
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        lines = md_content.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Skip empty lines
            if not stripped:
                i += 1
                continue

            # Horizontal rule — add a thin bottom border to the previous paragraph
            if stripped in ('---', '***', '___'):
                # Add a subtle divider via an empty paragraph with bottom border
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                pPr = p._element.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '4')
                bottom.set(qn('w:space'), '1')
                bottom.set(qn('w:color'), '999999')
                pBdr.append(bottom)
                pPr.append(pBdr)
                i += 1
                continue

            # H1 heading
            if stripped.startswith('# ') and not stripped.startswith('## '):
                text = stripped[2:].strip()
                p = doc.add_heading(level=1)
                self._add_formatted_runs(p, text)
                i += 1
                continue

            # H3 heading (check before H2 since ### starts with ##)
            if stripped.startswith('### '):
                text = stripped[4:].strip()
                p = doc.add_heading(level=3)
                self._add_formatted_runs(p, text)
                i += 1
                continue

            # H2 heading — uppercase to match PDF text-transform: uppercase
            if stripped.startswith('## '):
                text = stripped[3:].strip().upper()
                p = doc.add_heading(level=2)
                self._add_formatted_runs(p, text)
                i += 1
                continue

            # Bullet list item
            if stripped.startswith('- '):
                text = stripped[2:].strip()
                p = doc.add_paragraph(style='List Bullet')
                self._add_formatted_runs(p, text)
                i += 1
                continue

            # Plain paragraph (contact info, dates, body text)
            p = doc.add_paragraph()
            self._add_formatted_runs(p, stripped)
            i += 1

    def _add_formatted_runs(self, paragraph, text: str) -> None:
        """Parse inline markdown formatting and add runs to a paragraph.

        Handles **bold**, *italic*, and plain text segments.
        Processes the text left-to-right, splitting on formatting markers.

        Args:
            paragraph: python-docx Paragraph instance.
            text: Markdown text with potential inline formatting.
        """
        # Pattern matches **bold**, *italic*, or plain text segments
        # Process left-to-right: find the next formatting marker
        pos = 0
        length = len(text)

        while pos < length:
            # Look for the next ** or * marker
            bold_start = text.find('**', pos)
            italic_start = text.find('*', pos)

            # If italic_start points to a **, skip to real italic
            if italic_start != -1 and italic_start == bold_start:
                italic_start = text.find('*', bold_start + 2)
                # But check it's not another **
                if italic_start != -1 and italic_start + 1 < length and text[italic_start + 1] == '*':
                    italic_start = -1  # no standalone italic before next bold

            # Determine which comes first
            next_bold = bold_start if bold_start != -1 else length
            next_italic = italic_start if italic_start != -1 else length

            if next_bold <= next_italic and bold_start != -1:
                # Plain text before the bold marker
                if bold_start > pos:
                    paragraph.add_run(text[pos:bold_start])

                # Find closing **
                close = text.find('**', bold_start + 2)
                if close == -1:
                    # No closing marker — treat as plain text
                    paragraph.add_run(text[pos:])
                    break

                bold_text = text[bold_start + 2:close]
                run = paragraph.add_run(bold_text)
                run.bold = True
                pos = close + 2

            elif next_italic < next_bold and italic_start != -1:
                # Plain text before the italic marker
                if italic_start > pos:
                    paragraph.add_run(text[pos:italic_start])

                # Find closing *  (but not **)
                search_from = italic_start + 1
                close = -1
                while search_from < length:
                    candidate = text.find('*', search_from)
                    if candidate == -1:
                        break
                    # Make sure it's not a ** marker
                    if candidate + 1 < length and text[candidate + 1] == '*':
                        search_from = candidate + 2
                        continue
                    close = candidate
                    break

                if close == -1:
                    paragraph.add_run(text[pos:])
                    break

                italic_text = text[italic_start + 1:close]
                run = paragraph.add_run(italic_text)
                run.italic = True
                pos = close + 1

            else:
                # No more formatting markers — add remaining text
                paragraph.add_run(text[pos:])
                break

    def _sanitize_filename(self, name: str) -> str:
        """Sanitize a string for use in filenames."""
        sanitized = re.sub(r'[<>:"/\\|?*]', '', name)
        sanitized = sanitized.strip()
        return sanitized[:50]

    def _print_analysis(self, job: dict, analysis_result) -> None:
        """Print job analysis summary."""
        job_summary = analysis_result.job_summary
        match = analysis_result.match_assessment
        recs = analysis_result.resume_recommendations

        # Job info panel
        console.print(Panel(
            f"[bold]{job_summary.get('title', job.get('title', 'Unknown'))}[/bold] at "
            f"[cyan]{job_summary.get('company', job.get('company', 'Unknown'))}[/cyan]\n\n"
            f"[bold]Experience Required:[/bold] {job_summary.get('experience_required', 'Not specified')}",
            title="Job Summary"
        ))

        # Match score
        score = match.get("overall_score", "?")
        score_color = "green" if isinstance(score, int) and score >= 75 else "yellow" if isinstance(score, int) and score >= 50 else "red"
        console.print(f"\n[bold]Match Score:[/bold] [{score_color}]{score}/100[/{score_color}]")

        # Strengths
        strengths = match.get("strengths", [])
        if strengths:
            console.print("\n[bold green]Your Strengths:[/bold green]")
            for s in strengths[:4]:
                console.print(f"  - {s}")

        # Gaps
        gaps = match.get("gaps", [])
        if gaps:
            console.print("\n[bold yellow]Potential Gaps:[/bold yellow]")
            for g in gaps[:3]:
                console.print(f"  - {g}")

        # Domain connections
        domain_connections = match.get("domain_connections", [])
        if domain_connections:
            dc_lines = []
            for dc in domain_connections:
                conn_type = dc.get("connection_type", "unknown").replace("_", " ").title()
                dc_lines.append(
                    f"[bold]{dc.get('candidate_experience', '?')}[/bold]\n"
                    f"  → {dc.get('target_domain', '?')} [dim]({conn_type})[/dim]\n"
                    f"  [italic]{dc.get('reasoning', '')}[/italic]"
                )
            console.print(Panel(
                "\n\n".join(dc_lines),
                title="[bold magenta]Domain Connections[/bold magenta]",
                border_style="magenta",
            ))

        # Keywords to include
        keywords = recs.get("keywords_to_include", [])
        if keywords:
            console.print("\n[bold]Keywords to Include:[/bold]")
            console.print(f"  {', '.join(keywords[:8])}")

        console.print("\n[dim]Run 'scout resume <job_id>' to generate a customized resume[/dim]")
