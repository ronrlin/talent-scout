"""Job Researcher Agent - analyzes jobs and generates customized resumes."""

import json
import re
from datetime import datetime, timezone
from pathlib import Path

import anthropic
from docx import Document
from rich.console import Console
from rich.panel import Panel
from rich.progress import Progress, SpinnerColumn, TextColumn
from rich.table import Table

from config_loader import get_anthropic_api_key, get_all_location_slugs

console = Console()

JOB_ANALYSIS_PROMPT = """You are a job analysis expert helping a candidate understand a job posting and how well they match.

Analyze the job posting and the candidate's resume to provide:
1. Key requirements extraction
2. Match assessment
3. Gaps analysis
4. Customization recommendations

Return your analysis as JSON:
{
  "job_summary": {
    "title": "Job title",
    "company": "Company name",
    "key_responsibilities": ["top 5 responsibilities"],
    "required_skills": ["must-have skills"],
    "preferred_skills": ["nice-to-have skills"],
    "experience_required": "years and type of experience",
    "education_required": "education requirements"
  },
  "match_assessment": {
    "overall_score": 0-100,
    "strengths": ["candidate strengths that match this role"],
    "gaps": ["areas where candidate may be weak"],
    "transferable_skills": ["skills that translate well to this role"]
  },
  "resume_recommendations": {
    "skills_to_emphasize": ["skills from resume to highlight"],
    "experience_to_highlight": ["specific experiences to feature"],
    "keywords_to_include": ["ATS keywords from job posting to incorporate"],
    "sections_to_adjust": ["suggested section modifications"]
  },
  "cover_letter_points": ["key points to address in cover letter"],
  "interview_prep": ["topics to prepare for based on job requirements"]
}

Be specific and actionable. Focus on helping the candidate present themselves optimally for this role."""

RESUME_GENERATION_PROMPT = """You are an expert resume writer. Create a customized resume tailored to a specific job posting.

Guidelines:
- Maintain truthfulness - only use information from the base resume
- Reorder and emphasize relevant experience
- Incorporate keywords from the job posting naturally
- Keep to 1-2 pages maximum
- Use action verbs and quantified achievements
- Tailor the professional summary to the target role

Output the resume in clean Markdown format with clear sections:
- Contact info header
- Professional Summary (3-4 sentences tailored to this role)
- Professional Experience (reverse chronological)
- Education
- Additional relevant sections as needed

Global constraints:
- Avoid generic phrases (e.g., "proven track record," "results-driven," "dynamic leader").
- Do NOT overemphasize years of experience or education.
- Do NOT list tools or frameworks unless they were used in production systems I owned or led.
- Prefer outcomes, constraints, and decisions over responsibilities.

PROFESSIONAL EXPERIENCE
- Treat my current and most recent role as the primary anchor.
- For each role, include 3–6 bullets max.
- Each bullet should describe:
  - a concrete system, process, or organizational change I owned or led
  - the constraint or problem it addressed
  - the measurable or observable outcome
- At least one bullet per role should reflect people leadership or organizational design.
- Use action-oriented language, but avoid resume clichés.

EDUCATION
- Include degrees succinctly at the end without emphasis.

Formatting:
- Use clear section headers.
- Keep bullets concise but specific.
- Output only the resume in Markdown.

Do NOT invent experiences or skills not in the base resume. Only reorganize and reframe existing content."""

COVER_LETTER_PROMPT = """You are an expert cover letter writer. Create a compelling cover letter for a job application.

Guidelines:
- Professional and direct tone
- 3-4 paragraphs maximum
- Opening: Hook + why this company/role
- Middle: 2-3 key qualifications with specific examples from resume
- Closing: Call to action + enthusiasm

Output in clean Markdown format. Do NOT invent experiences not in the resume."""

COVER_LETTER_SPECIFICITY_PROMPT = """You are an expert editor focused on making cover letters highly specific and tailored.

Your task: Review the cover letter below and identify any sentence that could plausibly appear in a cover letter for a DIFFERENT company without modification.

For each such generic sentence, either:
1. REWRITE it to include specific details about the target company, role, or how the candidate's specific experience relates to THIS company's unique situation
2. REMOVE it entirely if it adds no specific value

Examples of GENERIC sentences to fix:
- "I am excited to apply for this position" → Remove or specify what about THIS role
- "My experience aligns well with your needs" → Name the specific experience and specific need
- "I would welcome the opportunity to discuss" → Too boilerplate, rewrite or cut
- "I am confident I can contribute to your team" → Generic confidence statement
- "Throughout my career, I have developed strong leadership skills" → Which career moments? What kind of leadership?

Examples of SPECIFIC sentences to KEEP:
- "At Tesla, I built an LLM-driven diagnostic system that reduced fleet troubleshooting time by 40%"
- "Your focus on multi-agent orchestration mirrors the distributed AI systems I architected at Waymo"
- "The Manager of AI & Analytics role's emphasis on Spark and distributed computing directly matches my work scaling data pipelines at Apple"

Rules:
- Maintain the same overall structure and flow
- Keep all specific company names, metrics, technologies, and concrete achievements
- The result should read as if it could ONLY be sent to this specific company
- Output ONLY the revised cover letter in Markdown (no explanations)"""

RESUME_DEFENSIBILITY_PROMPT = """You are a rigorous resume editor ensuring authenticity and defensibility.

Your task: Review the resume below and identify any content that:
1. Sounds too generic or could appear on anyone's resume
2. Unrealistically mirrors the job description (keyword stuffing or parroting requirements)
3. Makes claims that couldn't be defended in an interview
4. Uses inflated language that overstates actual responsibilities
5. Lists skills or technologies without evidence of real usage in the experience section
6. Lists a core skills section without adding much value to the overall impact of the resume

For each problematic element:
- REWRITE it to be specific, grounded, and defensible
- REMOVE it if it adds no real value or can't be substantiated
- TONE DOWN language that sounds like marketing copy

Red flags to fix:
- "Expert in X" without specific examples → Either add the example or use "experienced with"
- "Led transformation of..." without concrete details → Add what specifically changed
- Skills listed that don't appear in any job bullet → Remove or add supporting evidence
- Metrics that seem invented or rounded too neatly → Make more realistic or remove
- Buzzwords from the job posting that weren't in the original resume → Remove unless genuinely applicable
- Superlatives like "exceptional," "outstanding," "best-in-class" → Replace with factual descriptions
- Remove the Core Skills or Core Capabilities section if it does not provide an essential value to the role at hand

The goal: A hiring manager should be able to ask about ANY line on this resume and receive a concrete, honest answer. Nothing should require backpedaling or clarification in an interview.

Rules:
- Preserve the structure and formatting
- Keep all genuinely specific achievements, metrics, and examples
- Check that the resume presents work history in chronological order
- The candidate should be able to speak confidently to every single bullet point
- Output ONLY the revised resume in Markdown (no explanations)"""

RESUME_IMPROVE_PROMPT = """You are an expert technical recruiter and hiring manager with deep experience reviewing senior engineering, data, product, and technical leadership resumes.

Your task is to iteratively evaluate and improve a resume so that it is well-aligned with a specific job description while remaining credible, accurate, and professional.

Follow this process exactly:

STEP 1 — Analyze Job Description
- Identify the core responsibilities, required skills, preferred qualifications, seniority expectations, and success signals.
- Summarize the role's implicit evaluation criteria (what a hiring manager actually cares about).

STEP 2 — Evaluate Resume Against Job Description
- Assess alignment across the following dimensions:
  - Role fit and seniority
  - Technical depth and relevance
  - Scope, ownership, and impact
  - Leadership and cross-functional influence
  - Domain familiarity
- Explicitly identify:
  - Strong matches
  - Partial matches
  - Gaps or weak signals
  - Content that is off-target or distracting

STEP 3 — Identify Improvement Opportunities
- Propose concrete, high-impact opportunities to improve alignment.
- Focus on:
  - Reframing existing experience (do NOT invent experience)
  - Clarifying scope, ownership, or outcomes
  - Improving emphasis, ordering, or wording
  - Removing or de-emphasizing irrelevant content
- Avoid generic advice. Each recommendation must map directly to a job requirement.

STEP 4 — Iteratively Revise the Resume
- Apply the proposed improvements directly to the resume.
- Preserve factual accuracy and credibility at all times.
- Do not exaggerate titles, scope, or responsibilities.
- Prefer precision and clarity over buzzwords.
- After each revision pass:
  - Re-evaluate the updated resume against the job description.
  - Identify remaining gaps or misalignments.
  - Make additional targeted revisions if warranted.

STEP 5 — Stopping Criteria
- Stop iterating only when ALL of the following are true:
  - The resume is strongly aligned with the job description's core requirements
  - The narrative is coherent, senior-appropriate, and compelling
  - No material gaps remain that can be addressed without fabricating experience
  - The resume would be credible and competitive if reviewed by a hiring manager for this role

CONSTRAINTS:
- Do not invent experience.
- Do not misrepresent titles, timelines, or scope.
- Do not optimize for keyword stuffing.
- Optimize for clarity, credibility, and role fit.

OUTPUT FORMAT:
Return your response as valid JSON:
{
  "improvement_summary": [
    "bullet point explaining a key change made",
    "bullet point explaining why the resume is now better suited",
    "bullet point noting any remaining limitations"
  ],
  "revised_resume": "The complete revised resume in Markdown format"
}"""


class JobResearcherAgent:
    """Agent that analyzes jobs and generates customized application materials."""

    def __init__(self, config: dict):
        self.config = config
        self.client = anthropic.Anthropic(api_key=get_anthropic_api_key())
        self.data_dir = Path(__file__).parent.parent / "data"
        self.output_dir = Path(__file__).parent.parent / "output"
        self.input_dir = Path(__file__).parent.parent / "input"

        # Ensure output directories exist
        (self.output_dir / "resumes").mkdir(parents=True, exist_ok=True)
        (self.output_dir / "cover-letters").mkdir(parents=True, exist_ok=True)
        (self.output_dir / "analysis").mkdir(parents=True, exist_ok=True)

    def analyze_job(self, job_id: str) -> dict | None:
        """Analyze a job posting and match against user's profile."""
        # Find the job
        job = self._find_job(job_id)
        if not job:
            console.print(f"[red]Job not found: {job_id}[/red]")
            return None

        # Load base resume
        resume_text = self._load_base_resume()
        if not resume_text:
            console.print("[red]Could not load base resume[/red]")
            return None

        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            console=console,
        ) as progress:
            task = progress.add_task("Analyzing job requirements...", total=None)

            # Analyze with Claude
            analysis = self._analyze_job_match(job, resume_text)

            if not analysis:
                console.print("[red]Failed to analyze job[/red]")
                return None

            progress.update(task, description="Saving analysis...")

            # Save analysis
            self._save_analysis(job_id, job, analysis)

        # Print summary
        self._print_analysis(job, analysis)

        return analysis

    def generate_resume(self, job_id: str) -> str | None:
        """Generate a customized resume for a job."""
        # Find the job
        job = self._find_job(job_id)
        if not job:
            console.print(f"[red]Job not found: {job_id}[/red]")
            return None

        # Determine role lens
        role_lens = self._determine_role_lens(job)
        console.print(f"[dim]Role lens: {role_lens}[/dim]")

        # Load base resume
        resume_text = self._load_base_resume()
        if not resume_text:
            console.print("[red]Could not load base resume[/red]")
            return None

        # Load analysis if exists
        analysis = self._load_analysis(job_id)

        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            console=console,
        ) as progress:
            task = progress.add_task("Generating customized resume...", total=None)

            # Generate with Claude
            resume_md = self._generate_resume_content(job, resume_text, analysis, role_lens)

            if not resume_md:
                console.print("[red]Failed to generate resume[/red]")
                return None

            # Second pass: review for defensibility
            progress.update(task, description="Reviewing for defensibility...")
            resume_md = self._refine_resume_defensibility(resume_md, job, resume_text)

            progress.update(task, description="Saving resume...")

            # Save markdown
            company_name = self._sanitize_filename(job.get("company", "Unknown"))
            job_title = self._sanitize_filename(job.get("title", "Unknown"))
            md_path = self.output_dir / "resumes" / f"Ron Lin Resume - {company_name} - {job_title}.md"

            with open(md_path, "w") as f:
                f.write(resume_md)

            # Try to generate PDF
            progress.update(task, description="Generating PDF...")
            pdf_path = self._convert_to_pdf(md_path, "resume")

        console.print(f"[green]Resume saved to:[/green] {md_path}")
        if pdf_path:
            console.print(f"[green]PDF saved to:[/green] {pdf_path}")

        return str(md_path)

    def generate_cover_letter(self, job_id: str) -> str | None:
        """Generate a cover letter for a job."""
        # Find the job
        job = self._find_job(job_id)
        if not job:
            console.print(f"[red]Job not found: {job_id}[/red]")
            return None

        # Determine role lens
        role_lens = self._determine_role_lens(job)
        console.print(f"[dim]Role lens: {role_lens}[/dim]")

        # Load base resume
        resume_text = self._load_base_resume()
        if not resume_text:
            console.print("[red]Could not load base resume[/red]")
            return None

        # Load analysis if exists
        analysis = self._load_analysis(job_id)

        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            console=console,
        ) as progress:
            task = progress.add_task("Generating cover letter...", total=None)

            # Generate with Claude
            cover_letter_md = self._generate_cover_letter_content(job, resume_text, analysis, role_lens)

            if not cover_letter_md:
                console.print("[red]Failed to generate cover letter[/red]")
                return None

            # Second pass: refine for specificity
            progress.update(task, description="Refining for specificity...")
            cover_letter_md = self._refine_cover_letter_specificity(cover_letter_md, job)

            progress.update(task, description="Saving cover letter...")

            # Save markdown
            company_name = self._sanitize_filename(job.get("company", "Unknown"))
            job_title = self._sanitize_filename(job.get("title", "Unknown"))
            md_path = self.output_dir / "cover-letters" / f"Cover Letter - {company_name} - {job_title}.md"

            with open(md_path, "w") as f:
                f.write(cover_letter_md)

            # Try to generate PDF
            progress.update(task, description="Generating PDF...")
            pdf_path = self._convert_to_pdf(md_path, "cover-letter")

        console.print(f"[green]Cover letter saved to:[/green] {md_path}")
        if pdf_path:
            console.print(f"[green]PDF saved to:[/green] {pdf_path}")

        return str(md_path)

    def _find_job(self, job_id: str) -> dict | None:
        """Find a job by ID across all location files."""
        # Get all location slugs from config
        all_slugs = get_all_location_slugs(self.config)

        for slug in all_slugs:
            jobs_file = self.data_dir / f"jobs-{slug}.json"
            if not jobs_file.exists():
                continue

            with open(jobs_file) as f:
                data = json.load(f)

            for job in data.get("jobs", []):
                if job.get("id") == job_id:
                    job["_location"] = slug
                    return job

        return None

    def _get_role_lens_guidance(self, role_lens: str, doc_type: str) -> str:
        """Return role-lens specific guidance for resume or cover letter generation."""
        guidance = {
            "engineering": {
                "resume": """This is an ENGINEERING role. Emphasize:
- Technical systems architecture and ownership
- Code, infrastructure, and platform decisions
- Scaling engineering teams and establishing technical practices
- Production reliability, observability, and operational excellence
- AI/ML systems from experimentation to production deployment
- Technical mentorship and growing engineers""",
                "cover_letter": """This is an ENGINEERING role. Frame experience around:
- Systems you built or architected and their technical constraints
- Engineering team leadership and scaling
- Production operations and reliability outcomes
- Technical decision-making and trade-offs"""
            },
            "product": {
                "resume": """This is a PRODUCT role. Emphasize:
- Product strategy, vision, and roadmap ownership
- Customer outcomes and business metrics
- Cross-functional leadership with engineering, design, sales
- Data-driven decision making and experimentation
- Market analysis and competitive positioning
- Prioritization frameworks and trade-off decisions""",
                "cover_letter": """This is a PRODUCT role. Frame experience around:
- Products you shaped and the customer/business outcomes
- Strategic decisions about what to build and why
- Working with engineering teams to deliver product value
- Metrics, experimentation, and iteration"""
            },
            "program": {
                "resume": """This is a PROGRAM role. Emphasize:
- Cross-functional coordination and delivery execution
- Stakeholder management across engineering, product, leadership
- Process design, risk management, and dependency tracking
- Program-level metrics, reporting, and visibility
- Driving alignment and unblocking teams
- Launch coordination and operational readiness""",
                "cover_letter": """This is a PROGRAM role. Frame experience around:
- Complex programs you drove to completion
- Cross-functional coordination and stakeholder alignment
- Process improvements and delivery outcomes
- Risk identification and mitigation"""
            }
        }
        return guidance.get(role_lens, guidance["engineering"]).get(doc_type, "")

    def _determine_role_lens(self, job: dict) -> str:
        """Determine the role lens (engineering | product | program) based on job title and description.

        The role_lens shapes the language and emphasis in resumes and cover letters:
        - engineering: Focus on technical systems, architecture, code, team scaling
        - product: Focus on product strategy, roadmaps, customer outcomes, metrics
        - program: Focus on cross-functional coordination, delivery, process, stakeholder management
        """
        title = job.get("title", "").lower()
        department = job.get("department", "").lower()
        requirements = job.get("requirements_summary", "").lower()

        # Check for product indicators
        product_keywords = ["product manager", "product lead", "product director", "tpm", "technical product"]
        if any(kw in title for kw in product_keywords):
            return "product"

        # Check for program indicators
        program_keywords = ["program manager", "program lead", "program director", "tpm", "technical program"]
        if any(kw in title for kw in program_keywords):
            return "program"

        # Check for engineering indicators (default for ambiguous cases with eng focus)
        engineering_keywords = ["engineering manager", "engineer", "software", "data engineer", "analytics engineer",
                                "director of engineering", "vp engineering", "head of engineering", "staff engineer"]
        if any(kw in title for kw in engineering_keywords):
            return "engineering"

        # Secondary check on department
        if "product" in department:
            return "product"
        if "program" in department:
            return "program"
        if "engineering" in department or "data" in department:
            return "engineering"

        # Default to engineering for technical roles
        return "engineering"

    def _load_base_resume(self) -> str | None:
        """Load the base resume from Word doc."""
        resume_path = self.input_dir / "base-resume.docx"

        if not resume_path.exists():
            console.print(f"[yellow]Resume not found at {resume_path}[/yellow]")
            return None

        try:
            doc = Document(resume_path)

            # Extract text from paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            return "\n\n".join(text_parts)

        except Exception as e:
            console.print(f"[red]Error reading resume: {e}[/red]")
            return None

    def _analyze_job_match(self, job: dict, resume_text: str) -> dict | None:
        """Analyze how well the resume matches the job."""
        job_text = json.dumps(job, indent=2)

        response = self.client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4096,
            system=JOB_ANALYSIS_PROMPT,
            messages=[
                {
                    "role": "user",
                    "content": f"""Analyze this job posting and candidate resume:

## JOB POSTING:
{job_text}

## CANDIDATE RESUME:
{resume_text}

Provide a detailed match analysis and recommendations.""",
                }
            ],
        )

        return self._parse_json_response(response.content[0].text)

    def _generate_resume_content(self, job: dict, resume_text: str, analysis: dict | None, role_lens: str) -> str | None:
        """Generate customized resume content."""
        job_text = json.dumps(job, indent=2)
        analysis_text = json.dumps(analysis, indent=2) if analysis else "No prior analysis"

        # Role-lens specific guidance
        role_lens_guidance = self._get_role_lens_guidance(role_lens, "resume")

        response = self.client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4096,
            system=RESUME_GENERATION_PROMPT,
            messages=[
                {
                    "role": "user",
                    "content": f"""Create a customized resume for this job:

## TARGET JOB:
{job_text}

## ROLE LENS: {role_lens.upper()}
{role_lens_guidance}

## BASE RESUME (source material - use only this information):
{resume_text}

## ANALYSIS & RECOMMENDATIONS:
{analysis_text}

Ensure that the Professional Summary and the resume overall reflect the {role_lens} lens.

Generate a tailored resume in Markdown format.
.""",
                }
            ],
        )

        return response.content[0].text

    def _generate_cover_letter_content(self, job: dict, resume_text: str, analysis: dict | None, role_lens: str) -> str | None:
        """Generate cover letter content."""
        job_text = json.dumps(job, indent=2)
        analysis_text = json.dumps(analysis, indent=2) if analysis else "No prior analysis"

        # Role-lens specific guidance
        role_lens_guidance = self._get_role_lens_guidance(role_lens, "cover_letter")

        response = self.client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=2048,
            system=COVER_LETTER_PROMPT,
            messages=[
                {
                    "role": "user",
                    "content": f"""Create a cover letter for this job application:

## TARGET JOB:
{job_text}

## ROLE LENS: {role_lens.upper()}
{role_lens_guidance}

## CANDIDATE RESUME:
{resume_text}

## ANALYSIS (if available):
{analysis_text}

Generate a concise, high-signal cover letter in Markdown format.

The purpose of the letter is to demonstrate—through concrete examples—how my past decisions, systems, and leadership approach prepare me to succeed in the referenced role. The letter should read as authored by a senior {role_lens} leader, not as an expression of enthusiasm or a restatement of the job description.

Structural requirement:
- Begin with exactly ONE framing sentence that establishes the overarching problem space, leadership focus, or type of systems I have built or led.
- This opening sentence must be general (no company names, no job title, no praise for the employer) and should orient the reader to the examples that follow.
- After the opening sentence, immediately move into concrete experience.

Hard constraints:
- Do NOT use generic motivation or evaluation phrases (e.g., "I'm excited to apply," "drawn to your mission," "aligns well," "directly applicable," "ideal fit," "passionate about," "cutting-edge," "innovative").
- Do NOT paraphrase, summarize, or mirror the job description or company values. Assume the reader already knows them.
- Do NOT mention years of experience, education, or personal interest in the company.
- Avoid sentences that explicitly explain *why* an experience matches the role; show the experience and let the relevance be implicit.

Content requirements:
- Reference the job title naturally within a sentence (not as a heading or label).
- Write 2–3 short paragraphs total (excluding greeting and closing).
- Each paragraph must be anchored in a specific system, organizational change, or responsibility I owned or led.
- At least one paragraph should reflect people leadership or organizational design decisions.
- Avoid listing multiple technologies or skills in a single sentence.
- Apply the {role_lens} lens when selecting which experiences to highlight.

Style requirements:
- Vary paragraph structure and sentence rhythm; avoid repeating the same narrative pattern across paragraphs.
- Prefer concrete actions, constraints, and outcomes over abstractions or summaries.
- Use a confident, matter-of-fact professional tone (neither deferential nor sales-oriented).
- No header titled "Cover Letter."

Length:
- Approximately 200–300 words total.

Output only the cover letter in Markdown.""",
                }
            ],
        )

        return response.content[0].text

    def _refine_cover_letter_specificity(self, cover_letter: str, job: dict) -> str:
        """Second pass: review and rewrite generic sentences to be company-specific."""
        job_text = json.dumps(job, indent=2)

        response = self.client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=2048,
            system=COVER_LETTER_SPECIFICITY_PROMPT,
            messages=[
                {
                    "role": "user",
                    "content": f"""Review this cover letter and rewrite any generic sentences to be specific to this company and role.

## TARGET COMPANY/ROLE:
{job_text}

## COVER LETTER TO REFINE:
{cover_letter}

Output only the refined cover letter in Markdown.""",
                }
            ],
        )

        return response.content[0].text

    def _refine_resume_defensibility(self, resume: str, job: dict, base_resume: str) -> str:
        """Second pass: review resume for defensibility and remove generic/inflated content."""
        job_text = json.dumps(job, indent=2)

        response = self.client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4096,
            system=RESUME_DEFENSIBILITY_PROMPT,
            messages=[
                {
                    "role": "user",
                    "content": f"""Review this tailored resume for defensibility and authenticity.

## TARGET JOB (for context on what might be keyword-stuffed):
{job_text}

## ORIGINAL BASE RESUME (ground truth - what the candidate actually did):
{base_resume}

## TAILORED RESUME TO REVIEW:
{resume}

Ensure every claim is defensible and grounded in the original resume. Remove or tone down anything that sounds inflated or mirrors the job description too closely. Output only the refined resume in Markdown.""",
                }
            ],
        )

        return response.content[0].text

    def _save_analysis(self, job_id: str, job: dict, analysis: dict) -> None:
        """Save job analysis to file."""
        output_path = self.output_dir / "analysis" / f"{job_id}-analysis.json"

        data = {
            "job_id": job_id,
            "job": job,
            "analysis": analysis,
            "analyzed_at": datetime.now(timezone.utc).isoformat(),
        }

        with open(output_path, "w") as f:
            json.dump(data, f, indent=2)

    def _load_analysis(self, job_id: str) -> dict | None:
        """Load existing analysis if available."""
        analysis_path = self.output_dir / "analysis" / f"{job_id}-analysis.json"

        if analysis_path.exists():
            with open(analysis_path) as f:
                data = json.load(f)
                return data.get("analysis")

        return None

    def find_document_by_job_id(self, job_id: str, doc_type: str) -> Path | None:
        """Find an existing markdown document by job_id.

        Args:
            job_id: The job ID to search for
            doc_type: Either 'resume' or 'cover-letter'

        Returns:
            Path to the markdown file if found, None otherwise
        """
        # First, find the job to get company and title
        job = self._find_job(job_id)
        if not job:
            return None

        company_name = self._sanitize_filename(job.get("company", "Unknown"))
        job_title = self._sanitize_filename(job.get("title", "Unknown"))

        if doc_type == "resume":
            search_dir = self.output_dir / "resumes"
            pattern = f"Ron Lin Resume - {company_name} - {job_title}.md"
        else:
            search_dir = self.output_dir / "cover-letters"
            pattern = f"Cover Letter - {company_name} - {job_title}.md"

        # Try exact match first
        exact_path = search_dir / pattern
        if exact_path.exists():
            return exact_path

        # Fall back to searching for partial match (in case of filename truncation)
        if search_dir.exists():
            for md_file in search_dir.glob("*.md"):
                if company_name in md_file.name and job_title[:20] in md_file.name:
                    return md_file

        return None

    def regenerate_pdf(self, md_path: Path, doc_type: str) -> Path | None:
        """Regenerate PDF from an existing markdown file.

        Args:
            md_path: Path to the markdown file
            doc_type: Either 'resume' or 'cover-letter'

        Returns:
            Path to the generated PDF if successful, None otherwise
        """
        return self._convert_to_pdf(md_path, doc_type)

    def improve_resume(self, job_id: str) -> str | None:
        """Iteratively improve an existing resume to better align with a job.

        Args:
            job_id: The job ID to improve the resume for

        Returns:
            Path to the improved resume if successful, None otherwise
        """
        # Find the job
        job = self._find_job(job_id)
        if not job:
            console.print(f"[red]Job not found: {job_id}[/red]")
            return None

        # Find existing resume for this job
        resume_path = self.find_document_by_job_id(job_id, "resume")
        if not resume_path:
            console.print(f"[red]No existing resume found for job: {job_id}[/red]")
            console.print("[dim]Run 'scout resume <job_id>' first to generate a resume.[/dim]")
            return None

        # Load the current resume
        with open(resume_path) as f:
            current_resume = f.read()

        # Load base resume for reference (original facts)
        base_resume = self._load_base_resume()
        if not base_resume:
            console.print("[red]Could not load base resume[/red]")
            return None

        # Build job description context
        job_context = self._build_job_context(job)

        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            console=console,
        ) as progress:
            task = progress.add_task("Analyzing and improving resume...", total=None)

            # Call Claude to improve the resume
            response = self.client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=8000,
                system=RESUME_IMPROVE_PROMPT,
                messages=[
                    {
                        "role": "user",
                        "content": f"""Improve this resume to better align with the target job.

TARGET JOB:
{job_context}

CURRENT TAILORED RESUME:
{current_resume}

BASE RESUME (original facts - do not invent beyond this):
{base_resume}

Analyze the current resume against the job requirements, identify improvement opportunities, and iteratively revise until the resume is strongly aligned with the job while remaining credible and defensible."""
                    }
                ]
            )

            result_text = response.content[0].text
            result = self._parse_json_response(result_text)

            if not result:
                console.print("[red]Failed to parse improvement response[/red]")
                return None

            progress.update(task, description="Saving improved resume...")

            # Extract the improved resume
            improved_resume = result.get("revised_resume", "")
            improvement_summary = result.get("improvement_summary", [])

            if not improved_resume:
                console.print("[red]No improved resume in response[/red]")
                return None

            # Save the improved resume (overwrite the existing one)
            with open(resume_path, "w") as f:
                f.write(improved_resume)

            progress.update(task, description="Generating PDF...")

            # Regenerate PDF
            pdf_path = self._convert_to_pdf(resume_path, "resume")

        # Print improvement summary
        console.print(Panel(
            "\n".join(f"• {item}" for item in improvement_summary),
            title="[bold green]Improvement Summary[/bold green]",
            border_style="green"
        ))

        console.print(f"\n[green]Resume improved and saved to:[/green] {resume_path}")
        if pdf_path:
            console.print(f"[green]PDF saved to:[/green] {pdf_path}")

        return str(resume_path)

    def _build_job_context(self, job: dict) -> str:
        """Build a comprehensive job context string for prompts."""
        parts = []
        parts.append(f"Company: {job.get('company', 'Unknown')}")
        parts.append(f"Title: {job.get('title', 'Unknown')}")
        parts.append(f"Location: {job.get('location', 'Unknown')}")

        if job.get("requirements_summary"):
            parts.append(f"\nRequirements:\n{job.get('requirements_summary')}")

        if job.get("responsibilities_summary"):
            parts.append(f"\nResponsibilities:\n{job.get('responsibilities_summary')}")

        if job.get("match_notes"):
            parts.append(f"\nMatch Notes:\n{job.get('match_notes')}")

        # Try to load full job details from research file
        company_name = job.get("company", "").lower().replace(" ", "-")
        research_file = self.data_dir / "research" / f"{company_name}.json"
        if research_file.exists():
            try:
                with open(research_file) as f:
                    research = json.load(f)
                company_info = research.get("company", {})
                if company_info.get("description"):
                    parts.append(f"\nCompany Description:\n{company_info.get('description')}")
                if company_info.get("engineering_culture"):
                    parts.append(f"\nEngineering Culture:\n{company_info.get('engineering_culture')}")
            except Exception:
                pass

        return "\n".join(parts)

    def _convert_to_pdf(self, md_path: Path, doc_type: str) -> Path | None:
        """Convert markdown to PDF using weasyprint."""
        try:
            import markdown
            from weasyprint import HTML, CSS

            # Read markdown
            with open(md_path) as f:
                md_content = f.read()

            # Normalize bullet characters to standard markdown
            # LLM sometimes outputs • instead of - which markdown doesn't recognize as lists
            md_content = re.sub(r'^•\s*', '- ', md_content, flags=re.MULTILINE)

            # Convert to HTML
            html_content = markdown.markdown(md_content, extensions=['tables'])

            # Load template based on document type
            template_dir = Path(__file__).parent.parent / "templates"
            template_file = "resume.html" if doc_type == "resume" else "cover-letter.html"
            template_path = template_dir / template_file

            with open(template_path) as f:
                template = f.read()

            # Insert content into template
            full_html = template.replace("{{content}}", html_content)

            # Generate PDF
            pdf_path = md_path.with_suffix('.pdf')
            HTML(string=full_html).write_pdf(pdf_path)

            return pdf_path

        except ImportError as e:
            console.print(f"[yellow]PDF generation skipped (missing dependency): {e}[/yellow]")
            return None
        except Exception as e:
            console.print(f"[yellow]PDF generation failed: {e}[/yellow]")
            return None

    def _sanitize_filename(self, name: str) -> str:
        """Sanitize a string for use in filenames."""
        # Remove or replace invalid characters
        sanitized = re.sub(r'[<>:"/\\|?*]', '', name)
        sanitized = sanitized.strip()
        return sanitized[:50]  # Limit length

    def _parse_json_response(self, text: str) -> dict | None:
        """Parse Claude's JSON response."""
        try:
            if "```json" in text:
                text = text.split("```json")[1].split("```")[0]
            elif "```" in text:
                text = text.split("```")[1].split("```")[0]
            return json.loads(text.strip())
        except json.JSONDecodeError as e:
            console.print(f"[red]Error parsing response: {e}[/red]")
            return None

    def _print_analysis(self, job: dict, analysis: dict) -> None:
        """Print job analysis summary."""
        job_summary = analysis.get("job_summary", {})
        match = analysis.get("match_assessment", {})
        recs = analysis.get("resume_recommendations", {})

        # Job info panel
        console.print(Panel(
            f"[bold]{job_summary.get('title', job.get('title', 'Unknown'))}[/bold] at "
            f"[cyan]{job_summary.get('company', job.get('company', 'Unknown'))}[/cyan]\n\n"
            f"[bold]Experience Required:[/bold] {job_summary.get('experience_required', 'Not specified')}",
            title="Job Summary"
        ))

        # Match score
        score = match.get("overall_score", "?")
        score_color = "green" if score >= 75 else "yellow" if score >= 50 else "red"
        console.print(f"\n[bold]Match Score:[/bold] [{score_color}]{score}/100[/{score_color}]")

        # Strengths
        strengths = match.get("strengths", [])
        if strengths:
            console.print("\n[bold green]✓ Your Strengths:[/bold green]")
            for s in strengths[:4]:
                console.print(f"  • {s}")

        # Gaps
        gaps = match.get("gaps", [])
        if gaps:
            console.print("\n[bold yellow]⚠ Potential Gaps:[/bold yellow]")
            for g in gaps[:3]:
                console.print(f"  • {g}")

        # Keywords to include
        keywords = recs.get("keywords_to_include", [])
        if keywords:
            console.print("\n[bold]Keywords to Include:[/bold]")
            console.print(f"  {', '.join(keywords[:8])}")

        console.print("\n[dim]Run 'scout resume <job_id>' to generate a customized resume[/dim]")
