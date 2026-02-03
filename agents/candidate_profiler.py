"""Candidate Profiler Agent - manages candidate profile and preferences."""

import hashlib
import json
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from rich.console import Console
from rich.panel import Panel
from rich.progress import Progress, SpinnerColumn, TextColumn
from rich.table import Table

from data_store import DataStore
from .base_agent import BaseAgent

console = Console()

PROFILE_EXTRACTION_PROMPT = """You are an expert resume parser. Extract structured information from this resume.

Return your response as valid JSON matching this schema:
{
  "identity": {
    "name": "Full Name",
    "email": "email@example.com or null",
    "phone": "phone number or null",
    "linkedin": "LinkedIn URL or null",
    "location": "City, State or null"
  },
  "summary": "A structured 2-3 sentence summary of the candidate's professional profile",
  "experience": [
    {
      "company": "Company Name",
      "title": "Job Title",
      "start_date": "Month Year or Year",
      "end_date": "Month Year, Year, or Present",
      "highlights": ["key accomplishment 1", "key accomplishment 2"],
      "skills_used": ["skill1", "skill2"]
    }
  ],
  "skills": {
    "technical": ["programming languages, frameworks, tools"],
    "leadership": ["leadership and management skills"],
    "domains": ["industry domains, areas of expertise"]
  },
  "education": [
    {
      "institution": "University Name",
      "degree": "Degree Type",
      "field": "Field of Study",
      "year": "Graduation Year or null"
    }
  ]
}

Extract information accurately from the resume. Do not invent or assume information not present."""


class CandidateProfilerAgent(BaseAgent):
    """Agent that manages candidate profile and integrates learned preferences."""

    def __init__(self, config: dict):
        """Initialize the candidate profiler agent.

        Args:
            config: Configuration dictionary.
        """
        super().__init__(config)
        self.data_store = DataStore(config)

    def get_profile(self) -> dict | None:
        """Get the current candidate profile.

        Returns:
            Profile dictionary or None if not found.
        """
        return self.data_store.get_profile()

    def refresh_profile(self) -> dict | None:
        """Re-parse the base resume and update the profile.

        Returns:
            Updated profile dictionary or None if failed.
        """
        # Load base resume
        resume_text = self._load_base_resume()
        if not resume_text:
            console.print("[red]Could not load base resume[/red]")
            return None

        # Calculate hash of resume for change detection
        resume_hash = hashlib.sha256(resume_text.encode()).hexdigest()

        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            console=console,
        ) as progress:
            task = progress.add_task("Parsing resume...", total=None)

            # Extract structured profile from resume
            try:
                extracted = self.client.complete_json(
                    system=PROFILE_EXTRACTION_PROMPT,
                    user=f"Parse this resume and extract structured information:\n\n{resume_text}",
                    max_tokens=4096,
                )
            except ValueError as e:
                console.print(f"[red]Failed to parse resume: {e}[/red]")
                return None

            progress.update(task, description="Building profile...")

            # Load config preferences
            config_prefs = self.config.get("preferences", {})

            # Load learned preferences if available
            learned_prefs = self.data_store.get_learned_preferences()

            # Build complete profile
            profile = {
                "version": "1.0",
                "generated_at": datetime.now(timezone.utc).isoformat(),
                "source_resume_hash": resume_hash,
                "identity": extracted.get("identity", {}),
                "summary": extracted.get("summary", ""),
                "experience": extracted.get("experience", []),
                "skills": extracted.get("skills", {}),
                "education": extracted.get("education", []),
                "preferences": {
                    "target_roles": config_prefs.get("target_roles", []),
                    "locations": config_prefs.get("locations", []),
                    "include_remote": config_prefs.get("include_remote", True),
                },
                "learned_preferences": self._extract_learned_preferences(learned_prefs),
            }

            progress.update(task, description="Saving profile...")

            # Save profile
            self.data_store.save_profile(profile)

        console.print("[green]Profile updated successfully[/green]")
        return profile

    def view_profile(self) -> None:
        """Display the current profile in a formatted view."""
        profile = self.get_profile()

        if not profile:
            console.print("[yellow]No profile found. Run 'scout profile --refresh' to generate.[/yellow]")
            return

        # Identity section
        identity = profile.get("identity", {})
        name = identity.get("name", "Unknown")
        email = identity.get("email", "")
        location = identity.get("location", "")

        console.print(Panel(
            f"[bold]{name}[/bold]\n"
            f"[dim]{email}[/dim]\n"
            f"[dim]{location}[/dim]",
            title="Candidate Profile",
        ))

        # Summary
        summary = profile.get("summary", "")
        if summary:
            console.print(f"\n[bold]Summary:[/bold]\n{summary}")

        # Experience
        experience = profile.get("experience", [])
        if experience:
            console.print("\n[bold]Experience:[/bold]")
            for exp in experience[:5]:
                title = exp.get("title", "Unknown")
                company = exp.get("company", "Unknown")
                dates = f"{exp.get('start_date', '?')} - {exp.get('end_date', '?')}"
                console.print(f"  • {title} at [cyan]{company}[/cyan] [dim]({dates})[/dim]")

            if len(experience) > 5:
                console.print(f"  [dim]... and {len(experience) - 5} more positions[/dim]")

        # Skills
        skills = profile.get("skills", {})
        if skills:
            console.print("\n[bold]Skills:[/bold]")
            technical = skills.get("technical", [])
            if technical:
                console.print(f"  Technical: {', '.join(technical[:8])}")
            leadership = skills.get("leadership", [])
            if leadership:
                console.print(f"  Leadership: {', '.join(leadership[:5])}")
            domains = skills.get("domains", [])
            if domains:
                console.print(f"  Domains: {', '.join(domains[:5])}")

        # Preferences
        prefs = profile.get("preferences", {})
        if prefs:
            console.print("\n[bold]Target Preferences:[/bold]")
            roles = prefs.get("target_roles", [])
            if roles:
                console.print(f"  Roles: {', '.join(roles[:4])}")
            locations = prefs.get("locations", [])
            if locations:
                remote = " + Remote" if prefs.get("include_remote") else ""
                console.print(f"  Locations: {', '.join(locations)}{remote}")

        # Learned preferences
        learned = profile.get("learned_preferences", {})
        if learned:
            console.print("\n[bold]Learned Preferences:[/bold]")
            primary_titles = learned.get("primary_titles", [])
            if primary_titles:
                console.print(f"  Priority titles: {', '.join(primary_titles[:4])}")
            red_flags = learned.get("red_flag_keywords", [])
            if red_flags:
                console.print(f"  Avoiding: {', '.join(red_flags[:4])}")

        # Metadata
        generated = profile.get("generated_at", "Unknown")
        console.print(f"\n[dim]Profile generated: {generated}[/dim]")

    def _load_base_resume(self) -> str | None:
        """Load the base resume from Word doc."""
        resume_path = self.input_dir / "base-resume.docx"

        if not resume_path.exists():
            console.print(f"[yellow]Resume not found at {resume_path}[/yellow]")
            return None

        try:
            doc = Document(resume_path)

            # Extract text from paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            return "\n\n".join(text_parts)

        except Exception as e:
            console.print(f"[red]Error reading resume: {e}[/red]")
            return None

    def _extract_learned_preferences(self, learned_prefs: dict | None) -> dict:
        """Extract relevant learned preferences for the profile.

        Args:
            learned_prefs: Full learned preferences dict or None.

        Returns:
            Simplified preferences dict for profile.
        """
        if not learned_prefs:
            return {}

        targeting = learned_prefs.get("improved_targeting", {})
        scoring = learned_prefs.get("scoring_adjustments", {})

        return {
            "primary_titles": targeting.get("primary_titles", []),
            "titles_to_avoid": targeting.get("titles_to_avoid", []),
            "must_have_keywords": targeting.get("must_have_keywords", []),
            "red_flag_keywords": targeting.get("red_flag_keywords", []),
            "ideal_company_profile": targeting.get("ideal_company_profile", ""),
            "boost_factors": scoring.get("boost_factors", []),
            "penalty_factors": scoring.get("penalty_factors", []),
            "insights": learned_prefs.get("insights", ""),
        }

    def get_profile_summary(self) -> str:
        """Get a text summary of the profile for prompt injection.

        Returns:
            Summary string or empty if no profile.
        """
        profile = self.get_profile()
        if not profile:
            return ""

        parts = []

        # Identity
        identity = profile.get("identity", {})
        name = identity.get("name")
        if name:
            parts.append(f"Candidate: {name}")

        # Summary
        summary = profile.get("summary", "")
        if summary:
            parts.append(f"Background: {summary}")

        # Current/recent role
        experience = profile.get("experience", [])
        if experience:
            current = experience[0]
            parts.append(
                f"Current role: {current.get('title', '?')} at {current.get('company', '?')}"
            )

        # Key skills
        skills = profile.get("skills", {})
        technical = skills.get("technical", [])
        if technical:
            parts.append(f"Key skills: {', '.join(technical[:6])}")

        # Target preferences
        prefs = profile.get("preferences", {})
        roles = prefs.get("target_roles", [])
        if roles:
            parts.append(f"Target roles: {', '.join(roles[:3])}")

        return "\n".join(parts)
