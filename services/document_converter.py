"""Document converter - PDF and DOCX generation from markdown.

Extracted from agents/application_composer.py (_convert_to_pdf, _convert_to_docx,
_parse_markdown_to_docx, _add_formatted_runs).
"""

import logging
import re
from pathlib import Path

logger = logging.getLogger(__name__)


def convert_document(md_path: Path, doc_type: str, output_format: str = "pdf") -> dict[str, Path | None]:
    """Convert markdown to the requested output format(s).

    Args:
        md_path: Path to the source markdown file.
        doc_type: Type of document ("resume" or "cover-letter").
        output_format: "pdf", "docx", or "both".

    Returns:
        Dict with "pdf" and/or "docx" keys mapping to output Paths (or None if failed).
    """
    results = {}

    if output_format in ("pdf", "both"):
        results["pdf"] = convert_to_pdf(md_path, doc_type)

    if output_format in ("docx", "both"):
        results["docx"] = convert_to_docx(md_path, doc_type)

    return results


def convert_to_pdf(md_path: Path, doc_type: str) -> Path | None:
    """Convert markdown to PDF using weasyprint."""
    try:
        import markdown
        from weasyprint import HTML

        with open(md_path) as f:
            md_content = f.read()

        # Normalize bullet characters
        md_content = re.sub(r'^•\s*', '- ', md_content, flags=re.MULTILINE)

        html_content = markdown.markdown(md_content, extensions=['tables'])

        # Load template based on document type
        template_dir = Path(__file__).parent.parent / "templates"
        template_file = "resume.html" if doc_type == "resume" else "cover-letter.html"
        template_path = template_dir / template_file

        with open(template_path) as f:
            template = f.read()

        full_html = template.replace("{{content}}", html_content)

        pdf_path = md_path.with_suffix('.pdf')
        HTML(string=full_html).write_pdf(pdf_path)

        return pdf_path

    except ImportError as e:
        logger.warning("PDF generation skipped (missing dependency): %s", e)
        return None
    except Exception as e:
        logger.warning("PDF generation failed: %s", e)
        return None


def convert_to_docx(md_path: Path, doc_type: str) -> Path | None:
    """Convert markdown to DOCX using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        md_content = md_path.read_text()

        # Normalize bullet characters
        md_content = re.sub(r'^•\s*', '- ', md_content, flags=re.MULTILINE)

        doc = Document()

        # Page setup
        for section in doc.sections:
            section.top_margin = Inches(0.4)
            section.bottom_margin = Inches(0.4)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # Configure default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        style.paragraph_format.line_spacing = 1.3
        style.paragraph_format.space_before = Pt(3)
        style.paragraph_format.space_after = Pt(3)

        # H1
        h1_style = doc.styles['Heading 1']
        h1_style.font.name = 'Arial'
        h1_style.font.size = Pt(16)
        h1_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        h1_style.font.bold = True
        h1_style.paragraph_format.space_before = Pt(0)
        h1_style.paragraph_format.space_after = Pt(2)

        # H2
        h2_style = doc.styles['Heading 2']
        h2_style.font.name = 'Arial'
        h2_style.font.size = Pt(11)
        h2_style.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        h2_style.font.bold = True
        h2_style.paragraph_format.space_before = Pt(12)
        h2_style.paragraph_format.space_after = Pt(4)
        h2_pPr = h2_style.element.get_or_add_pPr()
        h2_pBdr = OxmlElement('w:pBdr')
        h2_bottom = OxmlElement('w:bottom')
        h2_bottom.set(qn('w:val'), 'single')
        h2_bottom.set(qn('w:sz'), '4')
        h2_bottom.set(qn('w:space'), '1')
        h2_bottom.set(qn('w:color'), '999999')
        h2_pBdr.append(h2_bottom)
        h2_pPr.append(h2_pBdr)

        # H3
        h3_style = doc.styles['Heading 3']
        h3_style.font.name = 'Arial'
        h3_style.font.size = Pt(10)
        h3_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        h3_style.font.bold = False
        h3_style.paragraph_format.space_before = Pt(12)
        h3_style.paragraph_format.space_after = Pt(2)

        # List bullet style
        if 'List Bullet' in doc.styles:
            lb_style = doc.styles['List Bullet']
            lb_style.font.name = 'Arial'
            lb_style.font.size = Pt(10)
            lb_style.paragraph_format.space_before = Pt(1)
            lb_style.paragraph_format.space_after = Pt(1)

        _parse_markdown_to_docx(doc, md_content, doc_type)

        docx_path = md_path.with_suffix('.docx')
        doc.save(str(docx_path))
        return docx_path

    except ImportError as e:
        logger.warning("DOCX generation skipped (missing dependency): %s", e)
        return None
    except Exception as e:
        logger.warning("DOCX generation failed: %s", e)
        return None


def _parse_markdown_to_docx(doc, md_content: str, doc_type: str) -> None:
    """Parse markdown content and build DOCX document elements."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    lines = md_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped in ('---', '***', '___'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '999999')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # H1
        if stripped.startswith('# ') and not stripped.startswith('## '):
            text = stripped[2:].strip()
            p = doc.add_heading(level=1)
            _add_formatted_runs(p, text)
            i += 1
            continue

        # H3 (check before H2)
        if stripped.startswith('### '):
            text = stripped[4:].strip()
            p = doc.add_heading(level=3)
            _add_formatted_runs(p, text)
            i += 1
            continue

        # H2 — uppercase
        if stripped.startswith('## '):
            text = stripped[3:].strip().upper()
            p = doc.add_heading(level=2)
            _add_formatted_runs(p, text)
            i += 1
            continue

        # Bullet
        if stripped.startswith('- '):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_runs(p, text)
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        _add_formatted_runs(p, stripped)
        i += 1


def _add_formatted_runs(paragraph, text: str) -> None:
    """Parse inline markdown formatting and add runs to a paragraph."""
    pos = 0
    length = len(text)

    while pos < length:
        bold_start = text.find('**', pos)
        italic_start = text.find('*', pos)

        if italic_start != -1 and italic_start == bold_start:
            italic_start = text.find('*', bold_start + 2)
            if italic_start != -1 and italic_start + 1 < length and text[italic_start + 1] == '*':
                italic_start = -1

        next_bold = bold_start if bold_start != -1 else length
        next_italic = italic_start if italic_start != -1 else length

        if next_bold <= next_italic and bold_start != -1:
            if bold_start > pos:
                paragraph.add_run(text[pos:bold_start])

            close = text.find('**', bold_start + 2)
            if close == -1:
                paragraph.add_run(text[pos:])
                break

            bold_text = text[bold_start + 2:close]
            run = paragraph.add_run(bold_text)
            run.bold = True
            pos = close + 2

        elif next_italic < next_bold and italic_start != -1:
            if italic_start > pos:
                paragraph.add_run(text[pos:italic_start])

            search_from = italic_start + 1
            close = -1
            while search_from < length:
                candidate = text.find('*', search_from)
                if candidate == -1:
                    break
                if candidate + 1 < length and text[candidate + 1] == '*':
                    search_from = candidate + 2
                    continue
                close = candidate
                break

            if close == -1:
                paragraph.add_run(text[pos:])
                break

            italic_text = text[italic_start + 1:close]
            run = paragraph.add_run(italic_text)
            run.italic = True
            pos = close + 1

        else:
            paragraph.add_run(text[pos:])
            break
